"""DeepDive — generate a one-page company research brief (.docx).

Uses DuckDuckGo for web research and python-docx for document generation.
Reads fit criteria from config.yaml so scoring is personalized.

Usage:
    startup-radar deepdive "Company Name"

Also callable from the Streamlit dashboard and the /deepdive Claude Code skill.
"""

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from startup_radar.config import AppConfig, load_config

REPORTS_DIR = Path.cwd() / "reports"
REPORTS_DIR.mkdir(exist_ok=True)


def _web_search(query: str, max_results: int = 8) -> list[dict]:
    try:
        from duckduckgo_search import DDGS

        return list(DDGS().text(query, max_results=max_results))
    except Exception:
        return []


def _search_company(name: str) -> dict:
    """Gather structured info about a company from web search."""
    info = {
        "name": name,
        "description": "",
        "founded": "",
        "hq": "",
        "funding_rounds": [],
        "total_raised": "",
        "valuation": "",
        "investors": [],
        "competitors": [],
        "founders": [],
        "employees": "",
        "website": "",
        "hiring_signals": [],
    }

    queries = [
        f"{name} startup funding raised investors",
        f"{name} startup founders team",
        f"{name} competitors market",
    ]

    all_snippets = []
    for q in queries:
        results = _web_search(q)
        for r in results:
            all_snippets.append(r.get("body", ""))
            if not info["website"] and r.get("href", ""):
                href = r["href"]
                if name.lower().replace(" ", "") in href.lower().replace(" ", ""):
                    info["website"] = href

    text = " ".join(all_snippets)

    first = _web_search(f"{name} startup", max_results=1)
    if first:
        info["description"] = first[0].get("body", "")[:300]

    amt = re.search(r"\$[\d,.]+\s*[BM]\b|\$[\d,.]+\s*(?:million|billion)", text, re.IGNORECASE)
    if amt:
        info["total_raised"] = amt.group(0).strip()

    for m in re.finditer(
        r"(Series\s+[A-F]\d?\+?|Seed|Pre-?Seed)\s*(?:round\s*)?(?:of\s*)?\$?([\d,.]+\s*[BM])?",
        text,
        re.IGNORECASE,
    ):
        round_name = m.group(1).strip()
        round_amt = m.group(2).strip() if m.group(2) else ""
        if round_amt:
            round_amt = f"${round_amt}"
        info["funding_rounds"].append({"round": round_name, "amount": round_amt})

    seen_rounds = set()
    deduped = []
    for r in info["funding_rounds"]:
        if r["round"].lower() not in seen_rounds:
            seen_rounds.add(r["round"].lower())
            deduped.append(r)
    info["funding_rounds"] = deduped

    investor_pattern = r"(?:led by|investors?\s+include|backed by|participation from)\s+([^.]+)"
    for m in re.finditer(investor_pattern, text, re.IGNORECASE):
        for inv in re.split(r",\s*|\s+and\s+", m.group(1)):
            inv = inv.strip().rstrip(".")
            if inv and 2 < len(inv) < 50 and inv not in info["investors"]:
                info["investors"].append(inv)

    loc = re.search(
        r"(?:based in|headquartered in)\s+([^,.\n]+(?:,\s*[A-Za-z. ]+)?)", text, re.IGNORECASE
    )
    if loc:
        info["hq"] = loc.group(1).strip()

    founder_pattern = r"(?:founded by|co-?founders?)\s+([^.]+)"
    for m in re.finditer(founder_pattern, text, re.IGNORECASE):
        for f in re.split(r",\s*|\s+and\s+", m.group(1)):
            f = f.strip()
            if f and 2 < len(f) < 60 and f not in info["founders"]:
                info["founders"].append(f)

    return info


def _score_company(info: dict, cfg: AppConfig) -> tuple[float, str, str]:
    """Score a company against the user's fit criteria. Returns (score, label, rationale)."""
    dd_cfg = cfg.deepdive
    factors = dd_cfg.fit_factors
    tier1_vcs = [v.lower() for v in dd_cfg.tier1_vcs]
    thresholds = dd_cfg.thresholds
    targets = cfg.targets

    weight_map = {"high": 1.5, "medium": 1.0, "low": 0.5}
    scores = {}

    ind_keywords = [k.lower() for k in targets.industries]
    desc = (info.get("description", "") + " " + info.get("name", "")).lower()
    if ind_keywords and any(k in desc for k in ind_keywords):
        scores["industry_match"] = 9
    elif not ind_keywords:
        scores["industry_match"] = 7
    else:
        scores["industry_match"] = 3

    rounds = info.get("funding_rounds", [])
    total = info.get("total_raised", "")
    if any("series" in r.get("round", "").lower() for r in rounds):
        scores["funding_stage"] = 9
    elif total and re.search(r"\d", total):
        scores["funding_stage"] = 6
    else:
        scores["funding_stage"] = 4

    locs = [loc.lower() for loc in targets.locations]
    hq = info.get("hq", "").lower()
    if locs and any(loc in hq for loc in locs):
        scores["location"] = 9
    elif "remote" in " ".join(locs):
        scores["location"] = 7
    else:
        scores["location"] = 4

    scores["role_fit_signals"] = 6
    for signal in info.get("hiring_signals", []):
        for role in targets.roles:
            if role.lower() in signal.lower():
                scores["role_fit_signals"] = 9
                break

    if info.get("founders"):
        scores["founder_pedigree"] = 7
    else:
        scores["founder_pedigree"] = 5

    investor_names = [i.lower() for i in info.get("investors", [])]
    has_tier1 = any(vc in inv for vc in tier1_vcs for inv in investor_names)
    scores["vc_tier"] = 9 if has_tier1 else 5

    total_weight = 0
    weighted_sum = 0
    for factor, raw in scores.items():
        w = weight_map.get(getattr(factors, factor, "medium"), 1.0)
        weighted_sum += raw * w
        total_weight += 10 * w

    final = round((weighted_sum / total_weight) * 10, 1) if total_weight else 5.0

    if final >= thresholds.strong:
        label = "STRONG FIT"
    elif final >= thresholds.moderate:
        label = "MODERATE FIT"
    else:
        label = "WEAK FIT"

    user_bg = cfg.user.background
    rationale_parts = []
    if has_tier1:
        rationale_parts.append("Backed by top-tier VCs")
    if scores.get("industry_match", 0) >= 8:
        rationale_parts.append("strong industry alignment")
    if scores.get("location", 0) >= 8:
        rationale_parts.append("located in your target area")
    if scores.get("role_fit_signals", 0) >= 8:
        rationale_parts.append("hiring for your target roles")
    if user_bg:
        rationale_parts.append(f"given your background ({user_bg[:80]})")
    rationale = (
        ". ".join(rationale_parts) + "."
        if rationale_parts
        else "Moderate alignment with your criteria."
    )

    return final, label, rationale


def _generate_docx(info: dict, score: float, label: str, rationale: str, cfg: AppConfig) -> Path:
    """Generate the .docx research brief."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    navy = RGBColor(0x1B, 0x3A, 0x6B)
    accent = RGBColor(0x2E, 0x75, 0xB6)

    h = doc.add_heading(info["name"], level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = navy
        run.font.size = Pt(22)

    if info.get("description"):
        p = doc.add_paragraph(info["description"])
        p.style.font.size = Pt(10)

    if info.get("website"):
        doc.add_paragraph(f"Website: {info['website']}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")

    def _section_header(text):
        h = doc.add_heading(text.upper(), level=2)
        for run in h.runs:
            run.font.color.rgb = accent
            run.font.size = Pt(12)

    _section_header("Company Overview")
    overview_items = [
        ("HQ", info.get("hq", "Not disclosed")),
        ("Founded", info.get("founded", "Not disclosed")),
        ("Total Raised", info.get("total_raised", "Not disclosed")),
        ("Employees", info.get("employees", "Not disclosed")),
    ]
    for label_text, val in overview_items:
        doc.add_paragraph(f"{label_text}: {val}")

    if info.get("founders"):
        _section_header("Founders & Team")
        for f in info["founders"]:
            doc.add_paragraph(f"- {f}")

    _section_header("Funding & Investors")
    if info.get("funding_rounds"):
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Round"
        hdr[1].text = "Amount"
        for r in info["funding_rounds"]:
            row = table.add_row().cells
            row[0].text = r.get("round", "")
            row[1].text = r.get("amount", "")

    tier1_vcs = [v.lower() for v in cfg.deepdive.tier1_vcs]
    if info.get("investors"):
        doc.add_paragraph("")
        p = doc.add_paragraph("Key Investors: ")
        for inv in info["investors"]:
            is_tier1 = any(vc in inv.lower() for vc in tier1_vcs)
            run = p.add_run(inv)
            if is_tier1:
                run.bold = True
            p.add_run(", ")

    if info.get("competitors"):
        _section_header("Competitors")
        for c in info["competitors"]:
            doc.add_paragraph(f"- {c}")

    _section_header("Fit Score")
    color = (
        RGBColor(0x1A, 0x7A, 0x4A)
        if "STRONG" in label
        else (RGBColor(0xB4, 0x53, 0x09) if "MODERATE" in label else RGBColor(0x99, 0x1B, 0x1B))
    )
    p = doc.add_paragraph()
    run = p.add_run(f"{score}/10 — {label}")
    run.font.size = Pt(14)
    run.font.color.rgb = color
    run.bold = True
    doc.add_paragraph(rationale)

    _section_header("Watch Out For")
    doc.add_paragraph("- Limited public information may indicate early stage or stealth mode.")
    if not info.get("investors"):
        doc.add_paragraph("- No investor information found — verify funding status independently.")
    if not info.get("hq"):
        doc.add_paragraph(
            "- Headquarters location not confirmed — may not match your target geography."
        )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated by Startup Radar \u00b7 {datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    safe_name = info["name"].strip().replace(" ", "")
    out_path = REPORTS_DIR / f"{safe_name}_Research_Brief.docx"
    doc.save(str(out_path))
    return out_path


def save_investors(company_name: str, investors: list[str]) -> Path:
    """Save investor list for warm-intro tier-2 lookup."""
    safe = company_name.strip().replace(" ", "")
    path = REPORTS_DIR / f"{safe}_investors.json"
    path.write_text(json.dumps(investors, indent=2))
    return path


def generate(company_name: str) -> Path:
    """Full pipeline: search, score, generate .docx, save investors."""
    cfg = load_config()
    print(f"Researching {company_name}...")
    info = _search_company(company_name)
    print("Scoring...")
    score, label, rationale = _score_company(info, cfg)
    print(f"Score: {score}/10 — {label}")
    print("Generating report...")
    path = _generate_docx(info, score, label, rationale, cfg)
    if info.get("investors"):
        save_investors(company_name, info["investors"])
    print(f"Report saved: {path}")
    return path
